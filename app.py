import io
import re

import streamlit as st
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


st.set_page_config(page_title="PDF → DOCX (mantieni struttura)", page_icon="📄")
st.title("📄 PDF → DOCX (mantieni struttura lista con immagini)")
st.caption("Carica un PDF. Il DOCX si scarica dopo la conversione (non va ricaricato).")

uploaded = st.file_uploader(
    "Carica un PDF (solo .pdf, NON Word)",
    type=["pdf"],
    help="Questa app converte PDF in Word. I file .docx non vanno caricati.",
)


def remove_table_borders(table):
    """Remove borders from a python-docx table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def repair_split_tokens(lines: list[str]) -> list[str]:
    """
    Fix common PDF text-splitting artifacts for this specific layout.
    Examples:
      "(X0026ZWS4B" + "-> 6"  -> "(X0026ZWS4B) -> 6"
      standalone "-> 6"      -> append to previous line
    """
    out = []
    i = 0
    while i < len(lines):
        a = lines[i].strip()

        # Case: "(X...." missing closing ")"
        if re.match(r"^\(X[0-9A-Z]{6,}$", a) and i + 1 < len(lines):
            b = lines[i + 1].strip()
            if re.match(r"^->\s*\d+$", b):
                out.append(f"{a}) {b}")
                i += 2
                continue
            out.append(f"{a})")
            i += 1
            continue

        # Case: line is only "-> N" -> attach to previous
        if re.match(r"^->\s*\d+$", a) and out:
            out[-1] = out[-1].rstrip() + " " + a
            i += 1
            continue

        out.append(a)
        i += 1

    return out


def normalize_text(t: str) -> str:
    """
    Clean and reorder text:
    - remove odd zero-width chars
    - remove empty lines
    - repair split tokens like (X...) and -> N
    - move '(CODE) -> N' lines to bottom of the product block
    """
    t = t.replace("\u2060", "").strip()
    lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
    lines = repair_split_tokens(lines)

    # Move code lines to bottom (so it resembles the PDF)
    code_lines = [ln for ln in lines if re.search(r"\(X[0-9A-Z]+\)\s*->\s*\d+", ln)]
    other = [ln for ln in lines if ln not in code_lines]
    return "\n".join(other + code_lines)


def extract_rows(page: fitz.Page):
    """
    Extract list rows from a page for PDFs like:
      [thumbnail image] [product text... (XCODE) -> qty]
    We use image blocks as anchors and collect nearby text blocks on the right.
    """
    data = page.get_text("dict")
    blocks = data.get("blocks", [])

    # Images: filter out tiny/irrelevant ones by size threshold
    img_blocks = [b for b in blocks if b.get("type") == 1 and b.get("size", 0) > 10000]
    img_blocks.sort(key=lambda b: b["bbox"][1])  # sort by top y

    text_blocks = [b for b in blocks if b.get("type") == 0]

    rows = []
    if not img_blocks:
        return rows

    for i, img in enumerate(img_blocks):
        y0 = img["bbox"][1]
        y1 = img["bbox"][3]

        # Determine vertical band for this "row"
        if i > 0:
            prev_bottom = img_blocks[i - 1]["bbox"][3]
            top = (prev_bottom + y0) / 2
        else:
            top = y0 - 8

        if i + 1 < len(img_blocks):
            next_top = img_blocks[i + 1]["bbox"][1]
            bottom = (y1 + next_top) / 2
        else:
            # Last row often needs more room (avoid cutting text at end of page)
            bottom = page.rect.height + 50

        # Collect text blocks to the right of the thumbnails in this band
        parts = []
        for tb in text_blocks:
            x0, ty0, x1, ty1 = tb["bbox"]

            # ignore left column (thumbnail side)
            if x0 < 105:
                continue

            # must intersect our vertical band
            if ty1 < top or ty0 > bottom:
                continue

            s = []
            for line in tb.get("lines", []):
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                if line_text.strip():
                    s.append(line_text)
            txt = "\n".join(s).strip()
            if txt:
                parts.append((tb["bbox"][1], txt))

        parts.sort(key=lambda x: x[0])
        full_text = normalize_text("\n".join([p for _, p in parts]))

        # Fallback: if almost empty, extend bottom a bit and retry
        if len(full_text) < 20:
            parts2 = []
            for tb in text_blocks:
                x0, ty0, x1, ty1 = tb["bbox"]
                if x0 < 105:
                    continue
                if ty1 < top or ty0 > (bottom + 120):
                    continue

                s = []
                for line in tb.get("lines", []):
                    line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                    if line_text.strip():
                        s.append(line_text)
                txt = "\n".join(s).strip()
                if txt:
                    parts2.append((tb["bbox"][1], txt))

            parts2.sort(key=lambda x: x[0])
            full_text = normalize_text("\n".join([p for _, p in parts2]))

        rows.append(
            {
                "image_bytes": img.get("image", b""),
                "text": full_text,
            }
        )

    return rows


def add_header_from_first_page(doc: Document, pdf: fitz.Document):
    """Try to rebuild a simple header (title/subtitle) from first page top text."""
    first = pdf[0]
    blocks = [b for b in first.get_text("dict").get("blocks", []) if b.get("type") == 0]

    header_blocks = [b for b in blocks if b["bbox"][1] < 190]
    header_blocks.sort(key=lambda b: b["bbox"][1])

    for idx, b in enumerate(header_blocks):
        lines = []
        for line in b.get("lines", []):
            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
            if line_text.strip():
                lines.append(line_text)
        t = "\n".join(lines).replace("\u2060", "").strip()
        if not t:
            continue

        p = doc.add_paragraph(t)

        if idx == 0:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(22)
        elif idx == 1:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16)
        else:
            for r in p.runs:
                r.font.size = Pt(11)

    doc.add_paragraph("")


def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    doc = Document()
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    add_header_from_first_page(doc, pdf)

    for page in pdf:
        for row in extract_rows(page):
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            # Column widths tuned for this layout
            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(5.5)

            remove_table_borders(table)

            # Image cell
            cell_img = table.cell(0, 0)
            p_img = cell_img.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if row["image_bytes"]:
                try:
                    p_img.add_run().add_picture(io.BytesIO(row["image_bytes"]), width=Inches(0.9))
                except Exception:
                    # If an image fails to embed, just skip it
                    pass

            # Text cell
            cell_txt = table.cell(0, 1)
            p_txt = cell_txt.paragraphs[0]

            if row["text"]:
                lines = row["text"].splitlines()
                for i, line in enumerate(lines):
                    run = p_txt.add_run(line)
                    run.font.size = Pt(11)
                    if i < len(lines) - 1:
                        p_txt.add_run("\n")

            doc.add_paragraph("")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


if uploaded:
    try:
        st.info("Conversione in corso…")
        docx_bytes = convert_pdf_to_docx(uploaded.read())
        st.success("Fatto ✅")

        st.download_button(
            "Scarica DOCX",
            data=docx_bytes,
            file_name=uploaded.name.rsplit(".", 1)[0] + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error(f"Errore: {e}")
        st.write("Se il PDF è una scansione (immagine), serve OCR e il risultato può variare.")
